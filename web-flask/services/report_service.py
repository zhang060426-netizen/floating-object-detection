from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches

from services.file_storage_service import resolve_object_path

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


_PLACEHOLDER = "不可用"


def _text(value, default=_PLACEHOLDER):
    if value is None or value == "":
        return default
    return str(value)


def _number(value, default=_PLACEHOLDER):
    if value is None or value == "":
        return default
    if isinstance(value, float):
        return f"{value:.4f}".rstrip("0").rstrip(".")
    return str(value)


def _percent(value, default=_PLACEHOLDER):
    if value is None or value == "":
        return default
    try:
        return f"{float(value) * 100:.2f}%"
    except (TypeError, ValueError):
        return str(value)


def _dict(value):
    return value if isinstance(value, dict) else {}


def _list(value):
    return value if isinstance(value, list) else []


def _filename_from_file_info(info):
    info = _dict(info)
    object_key = info.get("object_key")
    if object_key:
        return Path(str(object_key).replace(chr(92), "/")).name
    return None


def _safe_image_path(file_info):
    info = _dict(file_info)
    bucket = info.get("bucket")
    object_key = info.get("object_key")
    if not bucket or not object_key:
        return None, "文件信息缺失"
    try:
        path = resolve_object_path(bucket, object_key)
    except ValueError:
        return None, "文件路径未通过安全校验"
    if not path.exists() or not path.is_file():
        return None, "文件不存在"
    if path.suffix.lower() not in {".jpg", ".jpeg", ".png", ".webp"}:
        return None, "文件类型不支持嵌入"
    return path, None


def _add_key_values(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = _text(value)
    return table


def _add_missing_notes(document, notes):
    document.add_heading("缺失字段说明", level=2)
    if not notes:
        document.add_paragraph("未发现影响 MVP 报告生成的缺失字段。")
        return
    for note in notes:
        document.add_paragraph(note, style="List Bullet")


def _detection_status(summary, detections, detection_result):
    if summary.get("detection_status"):
        return summary.get("detection_status")
    if detection_result:
        return "detected" if detections else "no_detection"
    return "检测结果不可用"


def _target_count(summary, detections):
    for key in ("total_detections", "object_count"):
        if summary.get(key) is not None:
            return summary.get(key)
    return len(detections)


def _timing_text(timing, detection_result):
    if not timing:
        legacy = _dict(detection_result).get("timing_ms")
        if legacy is None:
            return "耗时信息不可用"
        return f"timing_ms: {_number(legacy)} ms"
    parts = []
    for key in ("total_api_ms", "inference_ms", "preprocess_ms", "model_load_ms", "postprocess_ms", "result_image_save_ms", "record_save_ms"):
        if timing.get(key) is not None:
            parts.append(f"{key}: {_number(timing.get(key))} ms")
    for key, value in timing.items():
        if key not in {"total_api_ms", "inference_ms", "preprocess_ms", "model_load_ms", "postprocess_ms", "result_image_save_ms", "record_save_ms"}:
            parts.append(f"{key}: {_text(value)}")
    return "; ".join(parts) if parts else "耗时信息不可用"


def _add_detection_table(document, detections):
    document.add_heading("检测目标表格", level=2)
    if not detections:
        document.add_paragraph("未检测到目标或检测目标列表为空。")
        return
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["序号", "类别", "置信度", "BBox(xyxy)", "面积(px)", "Track ID"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for index, detection in enumerate(detections, start=1):
        det = _dict(detection)
        cells = table.add_row().cells
        cells[0].text = str(det.get("object_index", index - 1))
        cells[1].text = _text(det.get("class_display_name") or det.get("class_name") or det.get("class_id"))
        cells[2].text = _percent(det.get("confidence"))
        bbox = det.get("bbox_xyxy") or det.get("bbox") or det.get("bbox_xywhn")
        cells[3].text = _text(bbox)
        cells[4].text = _number(det.get("area_px"))
        cells[5].text = _text(det.get("track_id"))


def _add_image_section(document, title, image_info, missing_notes):
    document.add_heading(title, level=2)
    path, reason = _safe_image_path(image_info)
    if not path:
        message = f"{title}未嵌入：{reason}。"
        document.add_paragraph(message)
        missing_notes.append(message)
        return
    try:
        document.add_picture(str(path), width=Inches(5.8))
        document.add_paragraph(f"文件：{path.name}")
    except Exception as exc:  # python-docx can reject uncommon/corrupt image encodings.
        message = f"{title}未嵌入：图片读取失败（{exc.__class__.__name__}）。"
        document.add_paragraph(message)
        missing_notes.append(message)


def _add_detection_result_summary(document, detection_result, summary, detections, timing):
    document.add_heading("detection_result.v1 摘要", level=2)
    if not detection_result:
        document.add_paragraph("检测结果不可用：记录未保存 detection_result。")
        return
    schema = _dict(detection_result).get("schema_version")
    document.add_paragraph(f"schema_version: {_text(schema)}")
    rows = [
        ("total_detections", _target_count(summary, detections)),
        ("has_detections", summary.get("has_detections", summary.get("has_detection"))),
        ("detection_status", _detection_status(summary, detections, detection_result)),
        ("max_confidence", _percent(summary.get("max_confidence"))),
        ("avg_confidence", _percent(summary.get("avg_confidence", summary.get("mean_confidence")))),
        ("class_counts", summary.get("class_counts")),
        ("timing", _timing_text(timing, detection_result)),
    ]
    _add_key_values(document, rows)


def build_detection_report_docx(record):
    """Build a single detection-record Word report in memory.

    The caller is responsible for loading ``record`` through detection_service.get_record
    so export permissions stay identical to record-detail permissions.
    """
    detection_result = record.get("detection_result")
    if not isinstance(detection_result, dict):
        detection_result = {}
    summary = _dict(detection_result.get("summary"))
    model = _dict(detection_result.get("model"))
    timing = _dict(detection_result.get("timing"))
    detections = _list(detection_result.get("detections"))
    missing_notes = []

    if not detection_result:
        missing_notes.append("detection_result 缺失或为空，检测结果不可用。")
    if not timing and detection_result.get("timing_ms") is None:
        missing_notes.append("timing 缺失，耗时信息不可用。")
    if not record.get("result_image"):
        missing_notes.append("result_image 缺失，结果图不存在或未生成。")

    document = Document()
    document.add_heading("水面漂浮物检测报告", level=0)

    filename = _filename_from_file_info(record.get("original_image"))
    if not filename:
        missing_notes.append("原始文件名不可用。")

    document.add_heading("记录信息", level=2)
    _add_key_values(
        document,
        [
            ("记录 ID", record.get("id") or record.get("record_id")),
            ("文件名", filename),
            ("检测时间", record.get("create_time")),
            ("模型信息", " / ".join(filter(None, [model.get("model_name"), model.get("model_id") or record.get("model_id"), model.get("base_model")])) or record.get("model_id")),
            ("检测状态", _detection_status(summary, detections, detection_result)),
            ("目标数量", _target_count(summary, detections) if detection_result else _PLACEHOLDER),
            ("置信度阈值", _percent(record.get("confidence_threshold") if record.get("confidence_threshold") is not None else model.get("confidence_threshold"))),
            ("最高置信度", _percent(summary.get("max_confidence"))),
            ("平均置信度", _percent(summary.get("avg_confidence", summary.get("mean_confidence")))),
            ("耗时信息", _timing_text(timing, detection_result)),
        ],
    )

    _add_detection_result_summary(document, detection_result, summary, detections, timing)
    _add_detection_table(document, detections)
    _add_image_section(document, "原图", record.get("original_image"), missing_notes)
    if record.get("result_image"):
        _add_image_section(document, "结果图", record.get("result_image"), missing_notes)
    else:
        document.add_heading("结果图", level=2)
        document.add_paragraph("结果图不存在或未生成。")
    _add_missing_notes(document, missing_notes)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
